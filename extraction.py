# -*- coding: utf-8 -*-
"""Extracción y normalización de documentos a texto plano.

Decisión de diseño (briefing §4): todo se normaliza a texto con marcadores de
página `[p. N]` antes de enviarlo a la API, incluidos los PDF nativos. Así las
citas del informe llevan localizador y el abogado las encuentra en segundos.

PDF escaneados sin capa de texto: NO se intenta OCR. Se envía lo que haya y la
fase 0 los rechaza con `extraccion_insuficiente`.
"""

import io


def extraer_pdf(datos: bytes) -> str:
    """PDF → texto con marcador [p. N] al inicio de cada página."""
    import pdfplumber

    paginas = []
    with pdfplumber.open(io.BytesIO(datos)) as pdf:
        for n, pagina in enumerate(pdf.pages, start=1):
            texto = pagina.extract_text() or ""
            paginas.append(f"[p. {n}]\n{texto}")
    return "\n\n".join(paginas)


def extraer_docx(datos: bytes) -> str:
    """DOCX → texto plano (párrafos y tablas). DOCX no expone paginación real,
    por lo que no se insertan marcadores de página."""
    import docx

    documento = docx.Document(io.BytesIO(datos))
    partes = []
    for parrafo in documento.paragraphs:
        if parrafo.text.strip():
            partes.append(parrafo.text)
    for tabla in documento.tables:
        for fila in tabla.rows:
            celdas = [celda.text.strip() for celda in fila.cells]
            if any(celdas):
                partes.append(" | ".join(celdas))
    return "\n".join(partes)


def extraer_texto_plano(datos: bytes) -> str:
    """TXT / MD → texto, con tolerancia de codificación."""
    for codificacion in ("utf-8", "latin-1"):
        try:
            return datos.decode(codificacion)
        except UnicodeDecodeError:
            continue
    return datos.decode("utf-8", errors="replace")


def extraer_fichero(nombre: str, datos: bytes) -> str:
    """Despacha por extensión. Lanza ValueError si el formato no está soportado."""
    extension = nombre.rsplit(".", 1)[-1].lower() if "." in nombre else ""
    if extension == "pdf":
        return extraer_pdf(datos)
    if extension == "docx":
        return extraer_docx(datos)
    if extension in ("txt", "md"):
        return extraer_texto_plano(datos)
    raise ValueError(f"Formato no soportado: {nombre}")


def normalizar_documentos(ficheros: list) -> str:
    """Concatena varios ficheros (contrato + order form + DPA + anexos) en un
    único texto, separados por una cabecera clara con el nombre de cada uno,
    para que la fase 0 pueda poblar `estructura_documental`.

    `ficheros` es una lista de tuplas (nombre, bytes).
    """
    bloques = []
    for nombre, datos in ficheros:
        texto = extraer_fichero(nombre, datos)
        bloques.append(
            f"===== DOCUMENTO: {nombre} =====\n\n{texto.strip()}"
        )
    return "\n\n\n".join(bloques)
